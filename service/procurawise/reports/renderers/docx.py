"""Fase 23 (ADR 0023) - the only module allowed to import python-docx
(module name `docx`), used exclusively for report_type=="rfp_document"
(spec S10: "Word y PDF")."""

from io import BytesIO

from docx import Document as DocxDocument

from procurawise.reports.render_types import ReportDocument


def render(document: ReportDocument) -> bytes:
    docx_document = DocxDocument()
    docx_document.add_heading(document.title, level=0)
    docx_document.add_heading(document.subtitle, level=1)
    for line in document.metadata_lines:
        paragraph = docx_document.add_paragraph(line)
        paragraph.style = docx_document.styles["Caption"]

    for section in document.sections:
        docx_document.add_heading(section.heading, level=2)
        for text in section.paragraphs:
            docx_document.add_paragraph(text)
        if section.table is not None:
            headers = section.table.headers
            rows = section.table.rows
            table = docx_document.add_table(rows=1, cols=len(headers))
            table.style = "Light Grid Accent 1"
            for cell, header in zip(table.rows[0].cells, headers, strict=True):
                cell.text = header
            for row in rows:
                cells = table.add_row().cells
                for cell, value in zip(cells, row, strict=True):
                    cell.text = str(value)

    buffer = BytesIO()
    docx_document.save(buffer)
    return buffer.getvalue()
